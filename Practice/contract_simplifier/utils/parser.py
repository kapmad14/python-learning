# utils/parser.py
"""
Parser utilities: PDF/DOCX/text extraction and OCR via OCR.Space cloud.
Uses per-page OCR (PyMuPDF) with caching and progress updates.
"""

import io
import logging
import os
import hashlib
from typing import Optional, List

import pdfplumber
import requests
from docx import Document
import streamlit as st

# PyMuPDF (fitz) used for rendering PDF pages to images
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

logger = logging.getLogger(__name__)

# --------- helpers ----------
def _sha256(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()

def _get_ocr_api_key() -> str:
    if hasattr(st, "secrets") and "OCR_SPACE_API_KEY" in st.secrets:
        return st.secrets["OCR_SPACE_API_KEY"]
    return os.environ.get("OCR_SPACE_API_KEY", "")

# ---------- OCR.Space request helper (safe) ----------
def ocr_space_request(file_bytes: bytes, filename: str, language: str = "eng", timeout: int = 120) -> str:
    """
    Send bytes to OCR.Space and return recognized text.
    Returns empty string on any failure and logs error.
    """
    OCR_KEY = _get_ocr_api_key()
    if not OCR_KEY:
        logger.warning("OCR_SPACE_API_KEY not set; OCR will be skipped and return empty text.")
        return ""

    url = "https://api.ocr.space/parse/image"
    payload = {"apikey": OCR_KEY, "language": language, "isOverlayRequired": False}
    files = {"file": (filename, file_bytes)}

    try:
        resp = requests.post(url, data=payload, files=files, timeout=timeout)
        resp.raise_for_status()
    except requests.RequestException as e:
        logger.exception("OCR.Space request failed: %s", e)
        return ""

    try:
        result = resp.json()
    except Exception as e:
        logger.exception("OCR.Space returned non-JSON response: %s", e)
        return ""

    if result.get("IsErroredOnProcessing", False):
        logger.error("OCR.Space processing error: %s", result.get("ErrorMessage"))
        return ""

    parsed_texts = []
    for pr in result.get("ParsedResults", []):
        parsed_texts.append(pr.get("ParsedText", ""))

    return "\n".join(parsed_texts).strip()

# ---------- Cached OCR (per-page) ----------
@st.cache_data(show_spinner=False)
def cached_ocr(page_hash: str, filename: str, file_bytes: bytes) -> str:
    """
    Run OCR via OCR.Space and cache result keyed by page_hash + filename.
    Note: page_hash is included for human-level clarity; the function args contribute to the cache key.
    """
    try:
        return ocr_space_request(file_bytes, filename)
    except Exception as e:
        logger.exception("cached_ocr failed: %s", e)
        return ""

# ---------- pdfplumber extraction (selectable text) ----------
def _extract_text_from_pdf_plumber_bytes(b: bytes) -> str:
    try:
        stream = io.BytesIO(b)
        with pdfplumber.open(stream) as pdf:
            pages = []
            for page in pdf.pages:
                try:
                    t = page.extract_text() or ""
                    pages.append(t)
                except Exception as e:
                    logger.exception("pdfplumber page.extract_text failed on a page: %s", e)
                    pages.append("")
        return "\n".join(pages).strip()
    except Exception as e:
        logger.exception("pdfplumber extraction failed: %s", e)
        return ""

# ---------- Public functions ----------

def extract_text_from_pdf(file) -> str:
    """
    Extract text from a PDF:
    - Prefer selectable text via pdfplumber.
    - If empty, fall back to per-page OCR via PyMuPDF (fitz) rendering + OCR.Space per page.
    Returns combined text string.
    """
    # Read bytes
    try:
        if hasattr(file, "getvalue"):
            b = file.getvalue()
        elif isinstance(file, (bytes, bytearray)):
            b = bytes(file)
        else:
            # assume it's a path
            with open(file, "rb") as fh:
                b = fh.read()
    except Exception as e:
        logger.exception("Failed to read PDF bytes: %s", e)
        return ""

    # First try pdfplumber (fast, no external calls)
    text = _extract_text_from_pdf_plumber_bytes(b)
    if text and text.strip():
        return text

    # If we reach here, fallback to per-page OCR
    # Try using PyMuPDF to render pages
    if fitz is None:
        # If fitz unavailable, fallback to single-call OCR on whole PDF bytes
        logger.warning("PyMuPDF (fitz) not available — using OCR.Space on whole PDF bytes (no per-page progress).")
        filename = getattr(file, "name", "file.pdf")
        return cached_ocr(_sha256(b), filename, b)

    try:
        doc = fitz.open(stream=b, filetype="pdf")
    except Exception as e:
        logger.exception("PyMuPDF failed to open PDF: %s", e)
        # fallback to single OCR
        filename = getattr(file, "name", "file.pdf")
        return cached_ocr(_sha256(b), filename, b)

    total = doc.page_count
    if total == 0:
        return ""

    # progress bar
    progress = st.progress(0.0)
    text_chunks: List[str] = []
    for i in range(total):
        try:
            page = doc.load_page(i)
            pix = page.get_pixmap(dpi=200)  # create raster image
            img_bytes = pix.tobytes("png")
            # compute page-specific hash to cache per page
            page_hash = _sha256(img_bytes)
            filename_page = f"{getattr(file, 'name', 'file')}_page_{i+1}.png"
            # call cached OCR for this page
            page_text = cached_ocr(page_hash, filename_page, img_bytes)
            text_chunks.append(page_text)
        except Exception as e:
            logger.exception("Failed OCR on page %s: %s", i + 1, e)
            text_chunks.append("")
        # update progress (use fraction)
        try:
            progress.progress((i + 1) / total)
        except Exception:
            # in some environments progress.progress may behave differently — ignore
            pass

    try:
        progress.empty()
    except Exception:
        pass

    return "\n".join(text_chunks).strip()

def extract_text_from_scanned_pdf(file) -> str:
    """
    Explicit scanned-PDF OCR helper — uses the same per-page OCR routine as extract_text_from_pdf.
    """
    return extract_text_from_pdf(file)

def extract_text_from_docx(file) -> str:
    """
    Extract text from a Word (.docx) file. Accepts bytes-like or file-like.
    """
    try:
        if hasattr(file, "getvalue"):
            b = file.getvalue()
            stream = io.BytesIO(b)
            doc = Document(stream)
        else:
            doc = Document(file)
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.exception("extract_text_from_docx failed: %s", e)
        return ""

def extract_text_from_image(file) -> str:
    """
    Extract text from an image (png/jpg) using cached OCR.
    """
    try:
        if hasattr(file, "getvalue"):
            b = file.getvalue()
            filename = getattr(file, "name", "image.png")
        elif isinstance(file, (bytes, bytearray)):
            b = bytes(file)
            filename = "image.png"
        else:
            with open(file, "rb") as fh:
                b = fh.read()
            filename = os.path.basename(file)
        page_hash = _sha256(b)
        return cached_ocr(page_hash, filename, b)
    except Exception as e:
        logger.exception("extract_text_from_image failed: %s", e)
        return ""
